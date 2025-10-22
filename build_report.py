# build_report.py
from __future__ import annotations
from pathlib import Path
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from unidecode import unidecode
import time
from datetime import datetime
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from pandas.api.types import is_datetime64_any_dtype as is_datetime
import os
# ===================== CONFIGURACIÓN =====================
# Carpeta base $Salida (equivalente a tu global en Stata)
if "RENDER" in os.environ:
    BASE_DIR = Path(__file__).resolve().parent
else:
    BASE_DIR = Path("/home/erick/projectos/stata/")
    
    
MES_NOMBRE = "Agosto"  # "$Mes" en Stata
MES_NUM = 8            # month(FechaPlan)==8
ANIO = 2025
API = True
# Estructura de carpetas (equivalente a "$Salida/..."):
DIR_INFO   = BASE_DIR / MES_NOMBRE / "Info"
DIR_APOYO  = BASE_DIR / MES_NOMBRE / "Apoyo"   # en pandas no guardamos; solo referencia conceptual
DIR_SALIDAS= BASE_DIR / MES_NOMBRE / "Salidas"
DIR_SALIDAS.mkdir(parents=True, exist_ok=True)

# Nombres de archivos (AJUSTA a los reales en tu carpeta Info)

try:
    from zoneinfo import ZoneInfo  # Py3.9+
except Exception:
    ZoneInfo = None
    
_MESES_ES = {
    1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril",
    5: "Mayo", 6: "Junio", 7: "Julio", 8: "Agosto",
    9: "Setiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"
}

def _nombre_mes_es(m: int, *, lower: bool = False) -> str:
    nombre = _MESES_ES.get(int(m), "")
    return nombre.lower() if lower else nombre

def compute_dia_actual_y_mes_pasado(
    *,
    tz: str = "America/Lima",
    base_mes_num: int | None = None,
    base_anio: int | None = None
) -> tuple[str, str]:
    """
    - dia_actual: fecha HOY formateada "dd de Mes de yyyy" (p.ej. 09 de Setiembre de 2025)
    - mes_pasado: mes anterior al 'base' en formato "mes de yyyy" en minúsculas (p.ej. "agosto de 2025")
        • si base_mes_num/base_anio se pasan, se usa ese mes como referencia
        • si no, se usa el mes actual como referencia
    """
    # now en Lima (o sistema si no hay zoneinfo)
    if ZoneInfo is not None:
        now = datetime.now(ZoneInfo(tz))
    else:
        now = datetime.now()

    # dia_actual (hoy)
    dia_actual = f"{now.day:02d} de {_nombre_mes_es(now.month)} de {now.year}"

    # referencia para mes_pasado (usa el mes del reporte si lo pasas)
    ref_m = base_mes_num if base_mes_num else now.month
    ref_y = base_anio if base_anio else now.year
    prev_m = 12 if ref_m == 1 else ref_m - 1
    prev_y = ref_y - 1 if ref_m == 1 else ref_y
    mes_pasado = f"{_nombre_mes_es(prev_m, lower=True)} de {prev_y}"

    return dia_actual, mes_pasado


def compute_dia_actual_y_mes_pasado_hoy_lima() -> tuple[str, str]:
    """dia_actual = 'dd de Mes de yyyy' (hoy, Lima) ; mes_pasado = '<mes-1> de yyyy' (en minúsculas)."""
    now = datetime.now(ZoneInfo("America/Lima")) if ZoneInfo else datetime.now()
    dia_actual = f"{now.day:02d} de {_nombre_mes_es(now.month)} de {now.year}"
    prev_m = 12 if now.month == 1 else now.month - 1
    prev_y = now.year - 1 if now.month == 1 else now.year
    mes_pasado = f"{_nombre_mes_es(prev_m, lower=True)} de {prev_y}"
    return dia_actual, mes_pasado

if not API:

    RUTAS_INFO = {
        "eventos_xlsx":   DIR_INFO / "Event_20250908201643.xlsx",           # sheet "Eventos"
        "vehiculos_xlsx": DIR_INFO / "2025-09-08_Vehículos PE - SUNASS.xlsx", # sheet "Vehículos"
        "ep_region_xlsx": BASE_DIR / "ep_region_macrorregion.xlsx", # sheet "Distribución_EP" (venía de $P_Salida)
        "direcciones_xlsx": DIR_INFO / "2025-09-08_Direcciones PE - SUNASS.xlsx"  # sheet "Direcciones"
    }
else:
    pass

# Salida final (un solo docx)
DOC_FINAL = DIR_SALIDAS / f"Informe_Final_{MES_NOMBRE}_{ANIO}.docx"


# ===================== UTILIDADES =====================
def set_style(doc: Document):
    """Estilo base del documento"""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    
def _iter_all_paragraphs(doc: Document):
    """Itera por todos los párrafos del body y dentro de tablas."""
    def _walk(container):
        if hasattr(container, "paragraphs"):
            for p in container.paragraphs:
                yield p
        if hasattr(container, "tables"):
            for t in container.tables:
                for row in t.rows:
                    for cell in row.cells:
                        yield from _walk(cell)
    yield from _walk(doc)
    
def replace_text_placeholder_anywhere_all(doc: Document, placeholder: str, value: str) -> int:
    """
    Reemplaza TODAS las apariciones de `placeholder` por `value` en el documento entero.
    Devuelve el número de reemplazos realizados.
    Nota: si el placeholder está dividido en varios runs, se reconstruye el párrafo con un único run.
    """
    if value is None:
        value = ""
    total = 0
    for p in _iter_all_paragraphs(doc):
        if placeholder in p.text:
            new_text = p.text.replace(placeholder, value)
            # limpiar todos los runs y dejar uno con el texto final
            for r in p.runs:
                r.text = ""
            if not p.runs:
                p.add_run(new_text)
            else:
                p.runs[0].text = new_text
            # contar cuántas veces se reemplazó en este párrafo
            # (aprox. por diferencia de longitudes)
            total += 1  # al menos 1 ocurrencia en este p
            # si quieres exactitud:
            # total += p.text.count(placeholder)  # usar antes del replace
    return total
    
def tune_word_table(table, doc, font_pt=7, equal_cols=True):
    # Ancho útil de la página (EMUs, es un entero)
    section = doc.sections[-1]
    usable_width = section.page_width - section.left_margin - section.right_margin

    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Fijar anchos por columna
    if equal_cols:
        n_cols = len(table.columns)
        if n_cols > 0:
            col_w = usable_width // n_cols   # ← entero (no float)
            for i in range(n_cols):
                table.columns[i].width = col_w
                for cell in table.columns[i].cells:
                    cell.width = col_w

    # Fuente tamaño 7
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_pt)

def add_paragraph(doc: Document, text: str, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p

def df_to_docx_table(doc: Document, df: pd.DataFrame, header_rename: dict | None = None):
    """Añade un DataFrame como tabla a un documento."""
    if df.empty:
        add_paragraph(doc, "(Sin datos)")
        return

    _df = df.copy()
    if header_rename:
        _df = _df.rename(columns=header_rename)

    rows, cols = _df.shape
    table = doc.add_table(rows=rows+1, cols=cols)
    table.style = "Light List Accent 1"

    # Header
    for j, col in enumerate(_df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for r in cell.paragraphs:
            for run in r.runs:
                run.bold = True

    # Body
    for i in range(rows):
        for j in range(cols):
            val = _df.iat[i, j]
            table.cell(i+1, j).text = "" if pd.isna(val) else str(val)

def parse_duration_like_stata(s: str) -> float:
    """
    Replica: split Duración by ":" into Time1,Time2,Time3
    DuracionAll = Time3 + Time2*60 + Time1*60*24
    """
    if pd.isna(s) or str(s).strip() == "":
        return np.nan
    parts = str(s).strip().split(":")
    parts = [p.strip() for p in parts]
    # Asegurar 3 componentes
    if len(parts) < 3:
        parts = (["0"] * (3 - len(parts))) + parts
    t1, t2, t3 = parts[-3], parts[-2], parts[-1]
    try:
        t1, t2, t3 = int(t1), int(t2), int(t3)
    except Exception:
        return np.nan
    return t3 + t2*60 + t1*60*24  # MISMA FÓRMULA QUE STATA


def build_All_strings(depositos: list[str]) -> tuple[str, str]:
    """
    Construye All y All_p como en Stata:
    - All = "(A), (B) y (C)"
    - All_p = "A, B y C" (sin paréntesis)
    """
    dep_fmt = [f"({d})" for d in depositos]
    if len(dep_fmt) == 0:
        return "", ""
    if len(dep_fmt) == 1:
        All = dep_fmt[0]
        All_p = All.replace("(", "").replace(")", "")
        return All, All_p
    # 2 o más
    All = ",  ".join(dep_fmt[:-1]) + f" y {dep_fmt[-1]}"
    All = All.replace("),(", "), (").replace(")y(", ") y (")
    All_p = All.replace("(", "").replace(")", "")
    return All, All_p


def normalize_vehicle_for_columns(v: str) -> str:
    """
    Emula: replace '-' by '_' y quitar espacios antes del reshape wide.
    Además, quitar tildes para evitar problemas de columnas en pandas/docx.
    """
    v2 = v.replace("-", "_").replace(" ", "")
    v2 = unidecode(v2)
    return v2


# ===================== PARTE A =====================
def cargar_y_procesar_eventos():
    df = pd.read_excel(RUTAS_INFO["eventos_xlsx"], sheet_name="Eventos")
    df.columns = [c.replace(" ", "") for c in df.columns]
    
    # Filtros
    df = df[~df["Vehículo"].astype(str).str.strip().eq("")]
    df = df[~df["Duración"].astype(str).str.strip().eq("")]

    # DuracionAll (misma fórmula Stata)
    df["DuracionAll"] = df["Duración"].map(parse_duration_like_stata)
    df = df[df["DuracionAll"].fillna(0) != 0]

    # Mes = month(FechaPlan) == 8
    # Asegurar que FechaPlan es datetime
    if not np.issubdtype(df["FechaPlan"].dtype, np.datetime64):
        df["FechaPlan"] = pd.to_datetime(df["FechaPlan"], errors="coerce", dayfirst=False)
    df = df[df["FechaPlan"].dt.month == MES_NUM]

    # Drop por texto
    df = df[~df["Depósito"].astype(str).str.contains("casa", case=False, na=False)]
    df = df[~df["Escenario"].astype(str).str.contains("Colegios", case=False, na=False)]

    df = df.sort_values("Vehículo", kind="stable").reset_index(drop=True)
    return df


def obtener_depositos(df_eventos: pd.DataFrame):
    depositos = (
        df_eventos["Depósito"]
        .dropna()
        .astype(str)
        .drop_duplicates()
        .tolist()
    )
    Cant = len(depositos)
    All, All_p = build_All_strings(depositos)
    return depositos, Cant, All, All_p


def seccion_tabla1_texto(doc: Document | None, All: str, *, mode: str = "docx", placeholder: str = "{{ texto_all }}"):
    texto = (
        f"En {MES_NOMBRE} de {ANIO}, la {All} enfrentaron interrupciones en el servicio y utilizaron la herramienta tecnológica para monitorear las operaciones de sus camiones cisterna, registrando tanto la ubicación como los horarios de las paradas realizadas durante el proceso de distribución y reabastecimiento de agua."
    )
    if mode == "docx":
        add_paragraph(doc, texto)
    else:  # "template"
        replace_text_placeholder_anywhere(doc, placeholder, texto)




def seccion_tabla1_por_deposito(doc: Document, df_eventos: pd.DataFrame, deposito: str):
    df_dep = df_eventos[df_eventos["Depósito"] == deposito].copy()
    cols = ["Vehículo","NombreEvento","Lat","Lng","Fecha","Hora","Duración"]
    cols = [c for c in cols if c in df_dep.columns]
    df_dep = df_dep[cols].sort_values("Fecha", kind="stable").reset_index(drop=True)

    # Cant_Vehi = # de vehículos únicos
    Cant_Vehi = df_dep["Vehículo"].nunique()

    # min_day / max_day (substr(Fecha,9,2) en Stata). Si es fecha, tomar el día; si es string, parse.
    if not np.issubdtype(df_dep["Fecha"].dtype, np.datetime64):
        df_dep["Fecha_dt"] = pd.to_datetime(df_dep["Fecha"], errors="coerce")
    else:
        df_dep["Fecha_dt"] = df_dep["Fecha"]

    min_day = int(df_dep["Fecha_dt"].dt.day.min()) if df_dep["Fecha_dt"].notna().any() else ""
    max_day = int(df_dep["Fecha_dt"].dt.day.max()) if df_dep["Fecha_dt"].notna().any() else ""

    texto = (
        f"{deposito}, utilizando {Cant_Vehi} camiones cisterna, empleó la herramienta tecnológica para el monitoreo de sus actividades "
        f"durante las interrupciones del servicio. Esto permitió registrar datos sobre la ubicación y los tiempos de las paradas realizadas "
        f"para la distribución o el reabastecimiento de agua del {min_day} al {max_day} de {MES_NOMBRE} de {ANIO}, "
        "tal como se detalla en la siguiente tabla:"
    )
    add_paragraph(doc, texto)

    header_map = {
        "NombreEvento": "Nombre evento",
        "Lat": "Latitud",
        "Lng": "Longitud"
    }
    df_to_docx_table(doc, df_dep[cols], header_rename=header_map)


def replace_text_placeholder_anywhere(doc: Document, placeholder: str, value: str):
    """Reemplaza texto simple en todo el documento (incluyendo celdas de tablas)."""
    if value is None:
        value = ""
    def _iter_paragraphs(parent):
        if hasattr(parent, "paragraphs"):
            for p in parent.paragraphs:
                yield p
        if hasattr(parent, "tables"):
            for t in parent.tables:
                for row in t.rows:
                    for cell in row.cells:
                        yield from _iter_paragraphs(cell)

    for p in _iter_paragraphs(doc):
        if placeholder in p.text:
            txt = p.text.replace(placeholder, value)
            for r in p.runs:
                r.text = ""
            if not p.runs:
                p.add_run(txt)
            else:
                p.runs[0].text = txt
            return True
    return False



def _insert_fast_table_after_paragraph(doc: Document, anchor_par, df, header_rename=None, table_style="Table Grid", font_pt=7, equal_cols=True):
    """
    Crea tabla rápida con df_to_docx_table_fast al final, la tunea, y la mueve
    justo DESPUÉS de anchor_par. Devuelve la tabla (objeto python-docx).
    """
    table = df_to_docx_table_fast(doc, df, header_rename=header_rename, table_style=table_style)
    tune_word_table(table, doc, font_pt=font_pt, equal_cols=equal_cols)
    # mover tabla debajo del anchor
    body = doc._element.body
    body.remove(table._tbl)                # quita del final
    anchor_par._p.addnext(table._tbl)      # inserta debajo del anchor
    return table



def _insert_paragraph_after_element(element, parent, text: str = ""):
    """
    Inserta un párrafo justo DESPUÉS de `element` (que puede ser paragraph._p o table._tbl).
    Devuelve el Paragraph recién creado.
    """
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
    new_p = OxmlElement("w:p")
    element.addnext(new_p)
    p2 = Paragraph(new_p, parent)
    if text:
        p2.add_run(text)
    return p2

def _insert_fast_table_after_element(doc, element, parent, df,
                                     header_rename=None,
                                     table_style="Table Grid",
                                     font_pt=7,
                                     equal_cols=True):
    """
    Crea tabla con df_to_docx_table_fast y la coloca INMEDIATAMENTE
    DESPUÉS de `element` (manteniendo el orden).
    """
    table = df_to_docx_table_fast(doc, df, header_rename=header_rename, table_style=table_style)
    tune_word_table(table, doc, font_pt=font_pt, equal_cols=equal_cols)
    # estaba al final; muévela justo después del element
    body = doc._element.body
    body.remove(table._tbl)
    element.addnext(table._tbl)
    return table

def render_series_text_plus_table_at_single_placeholder(
    doc: Document,
    placeholder: str,                       # ej: "{{ tabla_deposito }}"
    items: list[tuple[str, pd.DataFrame]],  # [(texto, df), ...] en orden
    header_rename: dict | None = None,
    table_style: str = "Table Grid",
    font_pt: int = 7,
    equal_cols: bool = True,
    show_separator: bool = False,
    separator_text: str = "—" * 40,
):
    """
    En el marcador único inserta, para CADA item:
      TEXTO -> TABLA -> (opcional SEP) -> (siguiente TEXTO) -> (siguiente TABLA) ...
    El ancla se va moviendo SIEMPRE al último elemento insertado (tabla) para garantizar alternancia.
    """
    anchor_par = find_placeholder_paragraph(doc, placeholder)
    if anchor_par is None:
        return False

    # Empezamos con el párrafo del marcador como "último elemento"
    last_el = anchor_par._p
    parent = anchor_par._parent

    for idx, (texto, df) in enumerate(items, start=1):
        # 1) TEXTO inmediatamente después del último elemento
        p_txt = _insert_paragraph_after_element(last_el, parent, texto)

        # 2) TABLA inmediatamente después del TEXTO
        if df is not None and not df.empty:
            tbl = _insert_fast_table_after_element(
                doc, p_txt._p, parent, df,
                header_rename=header_rename, table_style=table_style,
                font_pt=font_pt, equal_cols=equal_cols
            )
            last_el = tbl._tbl   # el nuevo ancla pasa a ser la TABLA (garantiza: texto -> tabla -> texto -> tabla)
        else:
            last_el = p_txt._p   # si no hay tabla, nos quedamos anclados al texto

        # 3) (Opcional) separador justo después del último elemento
        if show_separator and idx < len(items):
            p_sep = _insert_paragraph_after_element(last_el, parent, separator_text)
            last_el = p_sep._p

    # al final, eliminamos el párrafo del marcador original
    _remove_paragraph(anchor_par)
    return True





def seccion_tabla1_2_texto(doc: Document | None, All_p: str, *, mode: str = "docx", placeholder: str = "{{ texto_all_p }}") -> str | None:
    texto = (
        f"En {MES_NOMBRE} de {ANIO}, la herramienta tecnológica permitió registrar los datos de los camiones cisterna "
        f"y los puntos de abastecimiento de agua, información proporcionada por las empresas prestadoras {All_p} "
        f"a través del área usuaria."
    )
    if mode == "docx":
        # flujo sin plantilla: escribe un párrafo
        assert doc is not None, "doc requerido en mode='docx'"
        add_paragraph(doc, texto)
        return None
    elif mode == "template":
        # flujo con plantilla: reemplaza marcador
        assert doc is not None, "doc requerido en mode='template'"
        replace_text_placeholder_anywhere(doc, placeholder, texto)
        return texto  # por si quieres reusar el string
    else:
        return texto  # modo 'string' si lo quieres solo como retorno


def seccion_tabla1_3_por_deposito(doc: Document, df_eventos: pd.DataFrame, vehiculos_xlsx: Path, deposito: str):
    # Flotas
    df_flotas = pd.read_excel(vehiculos_xlsx, sheet_name="Vehículos")
    df_flotas.columns = [c.replace(" ", "") for c in df_flotas.columns]
    keep_cols = ["CódigodelVehículo","Descripción","Correo","Días","Flotas","Estado"]
    keep_cols = [c for c in keep_cols if c in df_flotas.columns]
    df_flotas = df_flotas[keep_cols].copy()

    # Códigos únicos reportados por depósito (renombrando Vehículo -> CódigodelVehículo)
    df_dep = df_eventos[df_eventos["Depósito"] == deposito].copy()
    df_dep = df_dep.rename(columns={"Vehículo":"CódigodelVehículo"})
    df_dep = df_dep[["CódigodelVehículo"]].drop_duplicates()

    # merge 1:1
    df_merge = df_dep.merge(df_flotas, on="CódigodelVehículo", how="inner")

    texto = (
        f"Durante el mes de {MES_NOMBRE} de {ANIO}, {deposito} registró en la herramienta tecnológica la siguiente "
        "información relacionada con sus camiones cisterna:"
    )
    add_paragraph(doc, texto)

    header_map = {"CódigodelVehículo":"Código del vehículo"}
    df_to_docx_table(doc, df_merge[keep_cols], header_rename=header_map)

def df_to_docx_table_fast(doc: Document, df, header_rename: dict | None = None, table_style: str = "Table Grid"):
    """
    Inserta un DataFrame como tabla en Word de forma eficiente,
    limpiando NaN/NaT/None -> "" y formateando fechas.
    """
    if header_rename:
        df = df.rename(columns=header_rename)

    # --- limpiar vacíos y formatear fechas ---
    _df = df.copy()
    for col in _df.columns:
        if is_datetime(_df[col]):
            _df[col] = _df[col].dt.strftime("%Y-%m-%d").fillna("")
        else:
            # convierte NaN/None/NaT y cadenas 'nan','NaN','None','NaT' a vacío
            _df[col] = (
                _df[col]
                .replace({np.nan: None})
                .astype(object)
                .where(~_df[col].astype(str).isin(["nan", "NaN", "None", "NaT"]), None)
                .fillna("")
                .astype(str)
            )

    cols = list(_df.columns)
    n_rows = len(_df) + 1
    n_cols = len(cols)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    if table_style:
        table.style = table_style

    # Encabezados
    hdr = table.rows[0].cells
    for j, col in enumerate(cols):
        hdr[j].text = str(col)

    # Filas
    for i, row in enumerate(_df.itertuples(index=False), start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            cells[j].text = "" if val is None else str(val)

    return table


def seccion_tabla1_4(
    doc: Document,
    ep_region_xlsx: Path,
    df_eventos: pd.DataFrame,
    direcciones_xlsx: Path,
    write_doc: bool = True,
    max_rows: int | None = None
):
    t0 = time.perf_counter()

    # 1) EP x Región
    df_ep = pd.read_excel(ep_region_xlsx, sheet_name="Distribución_EP",
                          dtype=str, engine="openpyxl")
    # Normaliza nombres quitando espacios (mantén acentos)
    df_ep.columns = [c.replace(" ", "") for c in df_ep.columns]

    if "EPS" in df_ep.columns:
        s = df_ep["EPS"].fillna("")
        s = (s.str.replace(" S.A.", "", regex=False)
               .str.replace(" S.R.L.", "", regex=False)
               .str.replace(".", "", regex=False))
        df_ep["EPS"] = s.replace({
            "EPS SEMAPACH SA": "SEMAPACH",
            "EPS EMAPA CAÑETE SA": "EMAPA CAÑETE",
            "EMAPA SAN MARTÍN": "EMAPA SAN MARTIN"
        })

    # 2) Diccionario Ubi por depósito
    depositos = (
        df_eventos[["Depósito"]]
        .drop_duplicates()
        .assign(EPS=lambda d: d["Depósito"].str.upper())
    )
    use_cols = ["EPS"] + (["REGIÓN"] if "REGIÓN" in df_ep.columns else [])
    df_ubi = depositos.merge(df_ep[use_cols], on="EPS", how="inner")

    if "REGIÓN" in df_ubi.columns:
        # Crear "Región" (mantén acento) para machar con df_dir
        df_ubi["Región"] = df_ubi["REGIÓN"].str.upper().replace({"SAN MARTÍN":"SAN MARTIN"})
    else:
        df_ubi["Región"] = ""

    # Provincias manuales por depósito (si aplica)
    df_ubi["Provincia"] = ""
    df_ubi.loc[df_ubi["Depósito"]=="EMAPA Cañete", "Provincia"] = "CAÑETE"
    df_ubi.loc[df_ubi["Depósito"]=="Sedachimbote", "Provincia"] = "CHIMBOTE"
    df_ubi.loc[df_ubi["Depósito"]=="Semapach", "Provincia"] = "CHINCHA"
    df_ubi.loc[df_ubi["Depósito"]=="EMAPA San Martin", "Provincia"] = "SAN MARTIN"
    df_ubi.loc[df_ubi["Depósito"]=="EPS BARRANCA", "Provincia"] = "BARRANCA"
    df_ubi = df_ubi[["Región","Provincia"]].drop_duplicates()

    t1 = time.perf_counter()
    print(f"[t] EP/Región y Ubi_dic listo en {t1 - t0:0.2f}s  (rows ubi={len(df_ubi)})")

    # 3) Direcciones (lee SOLO las columnas necesarias; conserva exactamente "Nombre de Dirección")
    use_cols_dir = ["Nombre de Dirección","Comuna","Provincia","Región","País","Lat","Lng"]
    df_dir = pd.read_excel(
        direcciones_xlsx,
        sheet_name="Direcciones",
        dtype=str,
        engine="openpyxl",
        usecols=use_cols_dir
    )

    # 4) Merge m:1 (filtrar por combinaciones válidas) → INNER para volver a ~2110
    if {"Región","Provincia"}.issubset(df_dir.columns) and not df_ubi.empty:
        pairs_ubi = df_ubi.drop_duplicates(["Región","Provincia"])[["Región","Provincia"]]
        n_all = len(df_dir)
        df_dir = df_dir.merge(pairs_ubi, on=["Región","Provincia"], how="inner")
        print(f"[dbg] filtrado por UBI: antes={n_all}, después={len(df_dir)}, pares={len(pairs_ubi)}")

    # Limita para pruebas
    if max_rows:
        df_dir = df_dir.head(max_rows)

    # Deduplicar SOLO si son filas idénticas (incluye el nombre para no colapsar distintas direcciones)
    subset_dd = [c for c in ["Nombre de Dirección","Comuna","Provincia","Región","País","Lat","Lng"] if c in df_dir.columns]
    df_dir = df_dir.drop_duplicates(subset=subset_dd)

    t2 = time.perf_counter()
    print(f"[t] Direcciones filtradas en {t2 - t1:0.2f}s  (rows dir={len(df_dir)})")

    if not write_doc:
        return df_dir

    # 5) Escritura al Word (rápida)
    texto = (
        f"Cabe mencionar que, los puntos de abastecimiento a camiones cisterna también fueron entregados por las "
        f"{df_eventos['Depósito'].nunique()} EPS a través del área usuaria y fueron ingresados a la herramienta:"
    )
    add_paragraph(doc, texto)

    # Renombres “bonitos” solo para la tabla
    header_map = {"Lat": "Latitud", "Lng": "Longitud"}
    table = df_to_docx_table_fast(doc, df_dir[use_cols_dir], header_rename=header_map)
    tune_word_table(table, doc, font_pt=7, equal_cols=True)

    t3 = time.perf_counter()
    print(f"[t] Escritura Word (tabla) {t3 - t2:0.2f}s")

    return df_dir

# ===================== PARTE B =====================
def seccion_tabla2_por_deposito(doc: Document, df_eventos: pd.DataFrame, deposito: str):
    df_dep = df_eventos[df_eventos["Depósito"] == deposito].copy()
    # collapse (sum) DuracionAll by(Fecha Vehículo)
    grp = (
        df_dep.groupby(["Fecha","Vehículo"], dropna=False)["DuracionAll"]
        .sum()
        .reset_index()
    )
    # normalizar nombres de vehículo para columnas
    grp["Vehiculo_norm"] = grp["Vehículo"].astype(str).map(normalize_vehicle_for_columns)

    # reshape wide
    tabla = grp.pivot_table(index="Fecha", columns="Vehiculo_norm", values="DuracionAll", aggfunc="sum", fill_value=0).reset_index()
    # Total rowtotal
    veh_cols = [c for c in tabla.columns if c not in ["Fecha"]]
    tabla["Total"] = tabla[veh_cols].sum(axis=1)

    # Orden por Fecha
    if not np.issubdtype(tabla["Fecha"].dtype, np.datetime64):
        tabla["Fecha_dt"] = pd.to_datetime(tabla["Fecha"], errors="coerce")
    else:
        tabla["Fecha_dt"] = tabla["Fecha"]
    tabla = tabla.sort_values("Fecha_dt", kind="stable").drop(columns=["Fecha_dt"])

    # Para cabeceras “bonitas”: reemplazar "_" por "-" (como tus labels en Stata)
    pretty_cols = ["Fecha"] + [c.replace("_","-") for c in veh_cols] + ["Total"]
    tabla.columns = pretty_cols

    texto = (
        f"A continuación, se presenta una tabla que resume, por día y por unidad, el tiempo total de monitoreo "
        f"(en las mismas unidades de 'DuracionAll' usadas por Stata) realizado por la {deposito} durante el mes de "
        f"{MES_NOMBRE} de {ANIO}:"
    )
    add_paragraph(doc, texto)
    df_to_docx_table(doc, tabla)
    
    
def build_tabla1_por_deposito(df_eventos: pd.DataFrame, deposito: str) -> pd.DataFrame:
    df_dep = df_eventos[df_eventos["Depósito"] == deposito].copy()
    cols = ["Vehículo","NombreEvento","Lat","Lng","Fecha","Hora","Duración"]
    cols = [c for c in cols if c in df_dep.columns]
    df_dep = df_dep[cols].sort_values("Fecha", kind="stable").reset_index(drop=True)
    return df_dep  # renombraremos encabezados al insertar

def build_tabla3_por_deposito(df_eventos: pd.DataFrame, vehiculos_xlsx: Path, deposito: str) -> pd.DataFrame:
    # Flotas
    df_flotas = pd.read_excel(vehiculos_xlsx, sheet_name="Vehículos")
    df_flotas.columns = [c.replace(" ", "") for c in df_flotas.columns]
    keep_cols = ["CódigodelVehículo","Descripción","Correo","Días","Flotas","Estado"]
    keep_cols = [c for c in keep_cols if c in df_flotas.columns]
    df_flotas = df_flotas[keep_cols].copy()

    df_dep = df_eventos[df_eventos["Depósito"] == deposito].copy()
    df_dep = df_dep.rename(columns={"Vehículo":"CódigodelVehículo"})
    df_dep = df_dep[["CódigodelVehículo"]].drop_duplicates()

    df_merge = df_dep.merge(df_flotas, on="CódigodelVehículo", how="inner")
    return df_merge  # encabezados bonitos al insertar

def build_tabla2_por_deposito(df_eventos: pd.DataFrame, deposito: str) -> pd.DataFrame:
    # Suma DuracionAll por día x vehículo y pivot
    df_dep = df_eventos[df_eventos["Depósito"] == deposito].copy()
    grp = (
        df_dep.groupby(["Fecha","Vehículo"], dropna=False)["DuracionAll"]
        .sum()
        .reset_index()
    )
    def normalize_vehicle_for_columns(v: str) -> str:
        from unidecode import unidecode
        v2 = v.replace("-", "_").replace(" ", "")
        v2 = unidecode(v2)
        return v2

    grp["Vehiculo_norm"] = grp["Vehículo"].astype(str).map(normalize_vehicle_for_columns)
    tabla = grp.pivot_table(index="Fecha", columns="Vehiculo_norm", values="DuracionAll", aggfunc="sum", fill_value=0).reset_index()

    # Total por fila
    veh_cols = [c for c in tabla.columns if c not in ["Fecha"]]
    tabla["Total"] = tabla[veh_cols].sum(axis=1)

    # Orden por fecha
    if not np.issubdtype(tabla["Fecha"].dtype, np.datetime64):
        tabla["Fecha_dt"] = pd.to_datetime(tabla["Fecha"], errors="coerce")
    else:
        tabla["Fecha_dt"] = tabla["Fecha"]
    tabla = tabla.sort_values("Fecha_dt", kind="stable").drop(columns=["Fecha_dt"])

    # Encabezados bonitos
    pretty_cols = ["Fecha"] + [c.replace("_","-") for c in veh_cols] + ["Total"]
    tabla.columns = pretty_cols
    return tabla

def build_tabla_grande(ep_region_xlsx: Path, df_eventos: pd.DataFrame, direcciones_xlsx: Path) -> pd.DataFrame:
    # Reusa tu lógica de seccion_tabla1_4 PERO retornando DF
    df_ep = pd.read_excel(ep_region_xlsx, sheet_name="Distribución_EP", dtype=str, engine="openpyxl")
    df_ep.columns = [c.replace(" ", "") for c in df_ep.columns]
    if "EPS" in df_ep.columns:
        s = df_ep["EPS"].fillna("")
        s = (s.str.replace(" S.A.", "", regex=False)
               .str.replace(" S.R.L.", "", regex=False)
               .str.replace(".", "", regex=False))
        df_ep["EPS"] = s.replace({
            "EPS SEMAPACH SA": "SEMAPACH",
            "EPS EMAPA CAÑETE SA": "EMAPA CAÑETE",
            "EMAPA SAN MARTÍN": "EMAPA SAN MARTIN"
        })
    depositos = df_eventos[["Depósito"]].drop_duplicates().assign(EPS=lambda d: d["Depósito"].str.upper())
    use_cols = ["EPS"] + (["REGIÓN"] if "REGIÓN" in df_ep.columns else [])
    df_ubi = depositos.merge(df_ep[use_cols], on="EPS", how="inner")
    if "REGIÓN" in df_ubi.columns:
        df_ubi["Región"] = df_ubi["REGIÓN"].str.upper().replace({"SAN MARTÍN":"SAN MARTIN"})
    else:
        df_ubi["Región"] = ""
    df_ubi["Provincia"] = ""
    df_ubi.loc[df_ubi["Depósito"]=="EMAPA Cañete", "Provincia"] = "CAÑETE"
    df_ubi.loc[df_ubi["Depósito"]=="Sedachimbote", "Provincia"] = "CHIMBOTE"
    df_ubi.loc[df_ubi["Depósito"]=="Semapach", "Provincia"] = "CHINCHA"
    df_ubi.loc[df_ubi["Depósito"]=="EMAPA San Martin", "Provincia"] = "SAN MARTIN"
    df_ubi.loc[df_ubi["Depósito"]=="EPS BARRANCA", "Provincia"] = "BARRANCA"
    df_ubi = df_ubi[["Región","Provincia"]].drop_duplicates()

    # Direcciones: columnas exactas (con espacio y tilde)
    use_cols_dir = ["Nombre de Dirección","Comuna","Provincia","Región","País","Lat","Lng"]
    df_dir = pd.read_excel(direcciones_xlsx, sheet_name="Direcciones", dtype=str, engine="openpyxl", usecols=use_cols_dir)

    # Filtrar por UBI (inner)
    if {"Región","Provincia"}.issubset(df_dir.columns) and not df_ubi.empty:
        pairs_ubi = df_ubi.drop_duplicates(["Región","Provincia"])[["Región","Provincia"]]
        df_dir = df_dir.merge(pairs_ubi, on=["Región","Provincia"], how="inner")

    # Deduplicación conservadora
    subset_dd = [c for c in ["Nombre de Dirección","Comuna","Provincia","Región","País","Lat","Lng"] if c in df_dir.columns]
    df_dir = df_dir.drop_duplicates(subset=subset_dd)

    return df_dir


def _remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None
    
    
def find_placeholder_paragraph(doc: Document, placeholder: str):
    def _iter_paragraphs(parent):
        if hasattr(parent, "paragraphs"):
            for p in parent.paragraphs:
                yield p
        if hasattr(parent, "tables"):
            for t in parent.tables:
                for row in t.rows:
                    for cell in row.cells:
                        yield from _iter_paragraphs(cell)

    for p in _iter_paragraphs(doc):
        if placeholder in p.text:
            return p
    return None


def place_df_at_placeholder(
    doc: Document,
    placeholder: str,
    df,
    header_rename=None,
    table_style="Table Grid",
    font_pt=7,
    equal_cols=True,
    before_text: str | None = None,   # 👈 nuevo
):
    p = find_placeholder_paragraph(doc, placeholder)
    if p is None:
        return None

    # 1) Si hay texto previo, insertarlo justo después del párrafo con el marcador
    anchor_el = p._p
    parent = p._parent
    if before_text:
        # crea un párrafo con el texto inmediatamente después del marcador
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph
        new_p = OxmlElement("w:p")
        anchor_el.addnext(new_p)
        p_txt = Paragraph(new_p, parent)
        p_txt.add_run(before_text)
        anchor_el = p_txt._p  # ahora anclamos después del texto

    # 2) Crear tabla rápida, tunearla y moverla justo después del texto (o del marcador si no hubo texto)
    table = df_to_docx_table_fast(doc, df, header_rename=header_rename, table_style=table_style)
    tune_word_table(table, doc, font_pt=font_pt, equal_cols=equal_cols)

    body = doc._element.body
    body.remove(table._tbl)          # estaba al final
    anchor_el.addnext(table._tbl)    # ahora queda justo después del texto/ancla

    # 3) Eliminar SOLO el párrafo del marcador original
    _remove_paragraph(p)

    return table


def render_tabla1_por_deposito_loop(doc, placeholder, df_eventos, depositos: list[str]):
    items = []
    for deposito in depositos:
        df_dep = df_eventos[df_eventos["Depósito"] == deposito].copy()
        cols = [c for c in ["Vehículo","NombreEvento","Lat","Lng","Fecha","Hora","Duración"] if c in df_dep.columns]
        df_dep = df_dep[cols].sort_values("Fecha", kind="stable").reset_index(drop=True)

        # métricas para el texto
        cant_vehi = df_dep["Vehículo"].nunique()
        if not np.issubdtype(df_dep["Fecha"].dtype, np.datetime64):
            df_dep["Fecha_dt"] = pd.to_datetime(df_dep["Fecha"], errors="coerce")
        else:
            df_dep["Fecha_dt"] = df_dep["Fecha"]
        min_day = int(df_dep["Fecha_dt"].dt.day.min()) if df_dep["Fecha_dt"].notna().any() else ""
        max_day = int(df_dep["Fecha_dt"].dt.day.max()) if df_dep["Fecha_dt"].notna().any() else ""

        texto = (
            f"{deposito}, utilizando {cant_vehi} camiones cisterna, empleó la herramienta tecnológica para el monitoreo de sus actividades "
            f"durante las interrupciones del servicio. Esto permitió registrar datos sobre la ubicación y los tiempos de las paradas realizadas "
            f"para la distribución o el reabastecimiento de agua del {min_day} al {max_day} de {MES_NOMBRE} de {ANIO}, "
            "tal como se detalla en la siguiente tabla:"
        )

        header_map = {"NombreEvento": "Nombre evento", "Lat": "Latitud", "Lng": "Longitud"}
        df_dep = df_dep.rename(columns=header_map)

        items.append((texto, df_dep))

    render_series_text_plus_table_at_single_placeholder(
        doc, placeholder, items,
        header_rename=None, table_style="Table Grid", font_pt=7, equal_cols=True,
        show_separator=False
    )
    
    
def render_tabla3_por_deposito_loop(doc, placeholder, df_eventos, vehiculos_xlsx: Path, depositos: list[str]):
    items = []
    header_map_3 = {"CódigodelVehículo": "Código del vehículo"}
    for deposito in depositos:
        texto = (
            f"Durante el mes de {MES_NOMBRE} de {ANIO}, {deposito} registró en la herramienta tecnológica la siguiente "
            "información relacionada con sus camiones cisterna:"
        )
        df3 = build_tabla3_por_deposito(df_eventos, vehiculos_xlsx, deposito).rename(columns=header_map_3)
        items.append((texto, df3))

    render_series_text_plus_table_at_single_placeholder(
        doc, placeholder, items,
        header_rename=None, table_style="Table Grid", font_pt=7, equal_cols=True,
        show_separator=False
    )

def render_tabla2_por_deposito_loop(doc, placeholder, df_eventos, depositos: list[str]):
    items = []
    for deposito in depositos:
        texto = (
            f"A continuación, se presenta una tabla que resume, por día y por unidad, el tiempo total de monitoreo "
            f"(en las mismas unidades de 'DuracionAll' usadas por Stata) realizado por la {deposito} durante el mes de "
            f"{MES_NOMBRE} de {ANIO}:"
        )
        df2 = build_tabla2_por_deposito(df_eventos, deposito)
        items.append((texto, df2))

    render_series_text_plus_table_at_single_placeholder(
        doc, placeholder, items,
        header_rename=None, table_style="Table Grid", font_pt=7, equal_cols=True,
        show_separator=False
    )
    
    

# ===================== ORQUESTA TODO =====================
def build_report():
    doc = Document()
    set_style(doc)

    # 1) Eventos procesados
    print("=== parte 1) Eventos procesados ===")
    df_eventos = cargar_y_procesar_eventos()

    # 2) Depósitos y textos All / All_p
    print("=== parte 2) Depósitos y textos All / All_p ===")
    depositos, Cant, All, All_p = obtener_depositos(df_eventos)

    # ---- Parte a) Texto general Tabla 1
    print("=== Parte a) Texto general Tabla 1 ===")
    seccion_tabla1_texto(doc, All)

    # ---- Parte a) por depósito (Tabla 1 por depósito)
    print("=== Parte a) por depósito (Tabla 1 por depósito) ===")
    for dep in depositos:
        print(dep)
        seccion_tabla1_por_deposito(doc, df_eventos, dep)

    # ---- Parte a)-2
    print("=== Parte a)-2 ===")
    seccion_tabla1_2_texto(doc, All_p)

    # ---- Parte a)-3 por depósito (merge con flotas)
    print("=== Parte a)-3 por depósito (merge con flotas) ===")
    for dep in depositos:
        seccion_tabla1_3_por_deposito(doc, df_eventos, RUTAS_INFO["vehiculos_xlsx"], dep)

    # ---- Parte a)-4 (Direcciones / Ubi)
    print("=== Parte a)-4 (Direcciones / Ubi) ===")
    df_dir_debug = seccion_tabla1_4(
        doc,
        RUTAS_INFO["ep_region_xlsx"],
        df_eventos,
        RUTAS_INFO["direcciones_xlsx"],
        write_doc=True,   # <- no genera tabla en docx
        max_rows=None       # <- limita filas para pruebas
    )
    print("Preview df_dir:", df_dir_debug.shape, df_dir_debug.head(3).to_dict(orient="records"))

    # ---- Parte b) por depósito (Pivot día x vehículo + Total)
    print("=== Parte b) por depósito (Pivot día x vehículo + Total) ===")
    for dep in depositos:
        print("dep2 :", dep)
        seccion_tabla2_por_deposito(doc, df_eventos, dep)

    # Guardar único Word
    print("===Guardar único Word ====")
    doc.save(DOC_FINAL)
    print(f"✅ Documento generado: {DOC_FINAL}")



def build_report_template_api(
    eventos_xlsx: Path,
    vehiculos_xlsx: Path,
    direcciones_xlsx: Path,
    ep_region_xlsx: Path,
    plantilla_docx: Path,
    salida_docx: Path,
    mes_nombre: str = MES_NOMBRE,
    mes_num: int = MES_NUM,
    anio: int = ANIO,
):
    # 1) Cargar y procesar
    df_eventos = pd.read_excel(eventos_xlsx, sheet_name="Eventos")
    df_eventos.columns = [c.replace(" ", "") for c in df_eventos.columns]
    df_eventos = df_eventos[~df_eventos["Vehículo"].astype(str).str.strip().eq("")]
    df_eventos = df_eventos[~df_eventos["Duración"].astype(str).str.strip().eq("")]
    df_eventos["DuracionAll"] = df_eventos["Duración"].map(parse_duration_like_stata)
    df_eventos = df_eventos[df_eventos["DuracionAll"].fillna(0) != 0]
    if not np.issubdtype(df_eventos["FechaPlan"].dtype, np.datetime64):
        df_eventos["FechaPlan"] = pd.to_datetime(df_eventos["FechaPlan"], errors="coerce", dayfirst=False)
    df_eventos = df_eventos[df_eventos["FechaPlan"].dt.month == mes_num]
    df_eventos = df_eventos[~df_eventos["Depósito"].astype(str).str.contains("casa", case=False, na=False)]
    df_eventos = df_eventos[~df_eventos["Escenario"].astype(str).str.contains("Colegios", case=False, na=False)]
    df_eventos = df_eventos.sort_values("Vehículo", kind="stable").reset_index(drop=True)

    # depósitos y textos
    depositos, Cant, All, All_p = obtener_depositos(df_eventos)
    deps = depositos                    # 👈 ya NO limitamos a 3

    # tablas auxiliares
    tabla_grande = build_tabla_grande(ep_region_xlsx, df_eventos, direcciones_xlsx)

    # 3) Abrir plantilla
    doc = Document(plantilla_docx)
    set_style(doc)

    dia_actual, mes_pasado = compute_dia_actual_y_mes_pasado_hoy_lima()
    replace_text_placeholder_anywhere_all(doc, "{{ dia_actual }}", dia_actual)
    replace_text_placeholder_anywhere_all(doc, "{{ mes_pasado }}", mes_pasado)

    # 4) Reemplazar textos generales si tienes marcadores en la plantilla
    # Uso (con plantilla):
    seccion_tabla1_texto(doc, All, mode="template", placeholder="{{ texto_all }}")
    seccion_tabla1_2_texto(doc, All_p, mode="template", placeholder="{{ texto_all_p }}")

    # 5) Tabla 1 por depósito (texto+tabla en bucle) en un solo marcador
    render_tabla1_por_deposito_loop(doc, "{{ tabla_deposito }}", df_eventos, deps)

    # 6) Tabla 3 por depósito (texto+tabla) en un solo marcador
    render_tabla3_por_deposito_loop(doc, "{{ tabla_3_por_deposito }}", df_eventos, vehiculos_xlsx, deps)

    # 7) Tabla grande (única)
    
    n_eps = df_eventos["Depósito"].nunique()
    texto_tabla_grande = (
        f"Cabe mencionar que, los puntos de abastecimiento a camiones cisterna también fueron entregados por las "
        f"{n_eps} EPS a través del área usuaria y fueron ingresados a la herramienta:"
    )
    
    header_map_grande = {"Lat": "Latitud", "Lng": "Longitud"}
    cols_grande = [c for c in ["Nombre de Dirección","Comuna","Provincia","Región","País","Lat","Lng"] if c in tabla_grande.columns]
    if cols_grande:
        place_df_at_placeholder(
            doc,
            "{{ tabla_grande }}",
            tabla_grande[cols_grande],
            header_rename=header_map_grande,
            font_pt=7,
            equal_cols=True,
            before_text=texto_tabla_grande,   # 👈 aquí el texto
        )

    # 8) Tabla 2 por depósito (texto+tabla) en un solo marcador
    render_tabla2_por_deposito_loop(doc, "{{ tabla_2_deposito }}", df_eventos, deps)

    # Guardar
    doc.save(salida_docx)
    print(f"✅ Documento generado con plantilla: {salida_docx}")
    return salida_docx



def build_report_template():
    # 1) Cargar y preparar
    df_eventos = cargar_y_procesar_eventos()
    depositos, Cant, All, All_p = obtener_depositos(df_eventos)

    # 2) Generar DataFrames (sin escribir en doc)
    # hasta 3 depósitos, en orden
    deps = depositos[:3]

    tablas_deposito = []
    for dep in deps:
        tablas_deposito.append(build_tabla1_por_deposito(df_eventos, dep))

    tablas_3_por_dep = []
    for dep in deps:
        tablas_3_por_dep.append(build_tabla3_por_deposito(df_eventos, RUTAS_INFO["vehiculos_xlsx"], dep))

    tablas_2_por_dep = []
    for dep in deps:
        tablas_2_por_dep.append(build_tabla2_por_deposito(df_eventos, dep))

    tabla_grande = build_tabla_grande(RUTAS_INFO["ep_region_xlsx"], df_eventos, RUTAS_INFO["direcciones_xlsx"])

    # 3) Abrir plantilla
    tpl_path = BASE_DIR / "plantilla_word_informe.docx"   # ← AJUSTA si está en otra carpeta
    doc = Document(tpl_path)
    set_style(doc)

    # 4) Colocar tablas en marcadores
    # {{ tabla_deposito_i }}
    header_map_dep = {"NombreEvento":"Nombre evento", "Lat":"Latitud", "Lng":"Longitud"}
    for i, df_dep in enumerate(tablas_deposito, start=1):
        placeholder = f"{{{{ tabla_deposito_{i} }}}}"
        cols = [c for c in ["Vehículo","NombreEvento","Lat","Lng","Fecha","Hora","Duración"] if c in df_dep.columns]
        place_df_at_placeholder(doc, placeholder, df_dep[cols], header_rename=header_map_dep, font_pt=7, equal_cols=True)

    # {{ tabla_3_por_deposito_i }}
    header_map_3 = {"CódigodelVehículo":"Código del vehículo"}
    for i, df3 in enumerate(tablas_3_por_dep, start=1):
        placeholder = f"{{{{ tabla_3_por_deposito_{i} }}}}"
        place_df_at_placeholder(doc, placeholder, df3, header_rename=header_map_3, font_pt=7, equal_cols=True)

    # {{ tabla_grande }}
    header_map_grande = {"Lat":"Latitud","Lng":"Longitud"}
    place_df_at_placeholder(
        doc,
        "{{ tabla_grande }}",
        tabla_grande[["Nombre de Dirección","Comuna","Provincia","Región","País","Lat","Lng"]],
        header_rename=header_map_grande,
        font_pt=7,
        equal_cols=True
    )

    # {{ tabla_2_deposito_i }}
    for i, df2 in enumerate(tablas_2_por_dep, start=1):
        placeholder = f"{{{{ tabla_2_deposito_{i} }}}}"
        place_df_at_placeholder(doc, placeholder, df2, header_rename=None, font_pt=7, equal_cols=True)

    # 5) Guardar
    doc.save(DOC_FINAL)
    print(f"✅ Documento generado con plantilla: {DOC_FINAL}")





if __name__ == "__main__":
    build_report_template()
